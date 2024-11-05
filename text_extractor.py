import os
from tkinter import Tk, Button, Text, filedialog, END, Scrollbar, RIGHT, Y, Menu
from PyPDF2 import PdfReader
from reportlab.lib.pagesizes import letter
from reportlab.pdfgen import canvas
import docx
from PIL import Image, ImageDraw, ImageFont
import pytesseract

def extract_text_from_pdf(pdf_path):
    with open(pdf_path, 'rb') as file:
        reader = PdfReader(file)
        text = ''
        for page in reader.pages:
            text += page.extract_text()
    return text

def extract_text_from_docx(docx_path):
    doc = docx.Document(docx_path)
    text = ''
    for paragraph in doc.paragraphs:
        text += paragraph.text + '\n'
    return text

def extract_text_from_image(image_path):
    image = Image.open(image_path)
    text = pytesseract.image_to_string(image)
    return text

def extract_text(file_path):
    _, file_extension = os.path.splitext(file_path)
    file_extension = file_extension.lower()

    if file_extension == '.pdf':
        return extract_text_from_pdf(file_path)
    elif file_extension == '.docx':
        return extract_text_from_docx(file_path)
    elif file_extension in ['.png', '.jpg', '.jpeg', '.tiff', '.bmp', '.gif']:
        return extract_text_from_image(file_path)
    else:
        return "Unsupported file format"

def upload_file():
    file_path = filedialog.askopenfilename()
    if file_path:
        text_box.delete(1.0, END)  # Clear previous text
        extracted_text = extract_text(file_path)
        text_box.insert(END, extracted_text)

def save_as_pdf(text, output_path):
    c = canvas.Canvas(output_path, pagesize=letter)
    width, height = letter
    text_object = c.beginText(40, height - 40)
    lines = text.split('\n')
    for line in lines:
        text_object.textLine(line)
    c.drawText(text_object)
    c.showPage()
    c.save()

def save_as_docx(text, output_path):
    doc = docx.Document()
    for line in text.split('\n'):
        doc.add_paragraph(line)
    doc.save(output_path)

def save_as_image(text, output_path):
    lines = text.split('\n')
    font = ImageFont.load_default()
    max_width = max(font.getsize(line)[0] for line in lines)
    total_height = sum(font.getsize(line)[1] for line in lines)
    image = Image.new('RGB', (max_width + 20, total_height + 20), 'white')
    draw = ImageDraw.Draw(image)
    y_text = 10
    for line in lines:
        draw.text((10, y_text), line, font=font, fill='black')
        y_text += font.getsize(line)[1]
    image.save(output_path)

def save_file():
    text = text_box.get(1.0, END)
    file_path = filedialog.asksaveasfilename(defaultextension=".txt", filetypes=[("PDF files", "*.pdf"), ("Word files", "*.docx"), ("Image files", "*.png"), ("All files", "*.*")])
    if file_path:
        _, file_extension = os.path.splitext(file_path)
        file_extension = file_extension.lower()
        if file_extension == '.pdf':
            save_as_pdf(text, file_path)
        elif file_extension == '.docx':
            save_as_docx(text, file_path)
        elif file_extension in ['.png', '.jpg', '.jpeg', '.tiff', '.bmp', '.gif']:
            save_as_image(text, file_path)
        else:
            text_box.insert(END, "\nUnsupported file format for saving.")

# Create the main window
root = Tk()
root.title("Text Extractor")
root.geometry("600x400")

# Create a button to upload files
upload_button = Button(root, text="Upload File", command=upload_file)
upload_button.pack(pady=10)

# Create a button to save files
save_button = Button(root, text="Save File", command=save_file)
save_button.pack(pady=10)

# Create a text box to display extracted text
text_box = Text(root, wrap='word', width=80, height=20)
text_box.pack(pady=10, padx=10)

# Add a scrollbar to the text box
scrollbar = Scrollbar(root, command=text_box.yview)
scrollbar.pack(side=RIGHT, fill=Y)
text_box.config(yscrollcommand=scrollbar.set)

# Run the application
root.mainloop()
